import re,requests,os,docx,time
from bs4 import BeautifulSoup
from selenium import webdriver
from docx import shared
from PIL import Image

# https://fy.58.com/chuzu/pn1/?key=%E9%98%9C%E9%98%B3%E7%A7%9F%E6%88%BF&final=1&PGTID=0d3090a7-0091-5e93-d284-a7804c9a5ecb&ClickID=2
def reqsuesturl(url):   #此函数是用来获取网页源代码，此函数调用频繁
    headers = {
        'authority': 'fy.58.com',
        'cache-control': 'max-age=0',
        'sec-ch-ua': '" Not;A Brand";v="99", "Google Chrome";v="91", "Chromium";v="91"',
        'sec-ch-ua-mobile': '?0',
        'upgrade-insecure-requests': '1',
        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.106 Safari/537.36',
        'accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
        'sec-fetch-site': 'none',
        'sec-fetch-mode': 'navigate',
        'sec-fetch-user': '?1',
        'sec-fetch-dest': 'document',
        'accept-language': 'zh-CN,zh;q=0.9',
        'cookie': 'f=n; commontopbar_new_city_info=2325%7C%E9%98%9C%E9%98%B3%7Cfy; commontopbar_ipcity=fy%7C%E9%98%9C%E9%98%B3%7C0; userid360_xml=66EA84D0F4D1501393B3E1DF4DB4CD47; time_create=1627019301185; id58=c5/nfGDSyyJvDmW6BLjYAg==; wmda_uuid=d3783c383f6dda711ebdf25afaf92214; wmda_new_uuid=1; wmda_session_id_11187958619315=1624427306959-877a17b0-542e-750a; wmda_visited_projects=%3B11187958619315; 58tj_uuid=743f50ad-b3b7-4fec-a74e-017966448b3f; new_uv=1; utm_source=; spm=; init_refer=; www58com="UserID=80415072209926&UserName=92n204kcx"; 58cooper="userid=80415072209926&username=92n204kcx"; 58uname=92n204kcx; PPU=UID=80415072209926&UN=92n204kcx&TT=34f99566f6245941e52df7a1076b9630&PBODY=iXhwpu3-7GgDAgfX74EaYcg4_4bw6wn7nhBWthTSEA6ubKKmRBCBbn7xdVo7-Pca9pl0l4Nx207PREjSqBdY82A6alLvdHl9rKg0q3Voq_GdteiAgJml9qjcIx4xHvNYMtaBcyvIm6Dk-vrALEBotb37T8m6zNIQVzZ8wkBZMKA&VER=1; new_session=0; ppStore_fingerprint=undefined%EF%BC%BF1624427323728; als=0; xxzl_deviceid=rSEDMhAhOunDOxMkQgKZtLbpvtmcEVJLdlvNiSi%2FJzENsSDDLMZ%2B%2BqL5J5w96qu8; xxzl_smartid=a74b90baf428fda6143f1f98dc453e4e',
    }

    params = (
        ('key', '\u961C\u9633\u79DF\u623F'),
        ('final', '1'),
        ('PGTID', '0d3090a7-0091-513b-f10d-8f814462c124'),
        ('ClickID', '3'),
    )

    r = requests.get(url, headers=headers,params=params)
    return r.text

def firsturl(url,list): #从一级网页中提取出二级网页
    html = reqsuesturl(url)
    soup = BeautifulSoup(html,'lxml')
    list1 = soup.find_all(class_='des')
    for i in list1:
        str1 = re.findall(r'<a class="strongbox" href=".*?"',str(i))[0]
        second_url = str1.split('href=')[1][1:-1]
        list.append(second_url)

def secondurl(url): #从二级网页中提取信息并保存在word文档里
    html = reqsuesturl(url)
    soup = BeautifulSoup(html, 'lxml')  #进行bs4解析
    try:
        message1 = (soup.find(class_='c_333 f20 strongbox').string).replace(' ', '_')  #如果正常运行说明没有遇到身份验证（反爬虫机制）
        # message1 = str(soup.find(class_='c_333 f20 strongbox')).split('>')[1][:-4]
        # 这里是标题的两种方法
    except:
        #异常说明了58同城的反爬机制上线！
        print("\033[31m这里进行了身份信息验证！\033[0m")
        print("\033[31m这里进行了身份信息验证！\033[0m")
        print("\033[31m这里进行了身份信息验证！\033[0m")
        print("\033[31m这里进行了身份信息验证！\033[0m")
        print("\033[31m这里进行了身份信息验证！\033[0m")
        print("\033[31m这里进行了身份信息验证！\033[0m")
        print("\033[31m这里进行了身份信息验证！\033[0m")
        print("\033[31m这里进行了身份信息验证！\033[0m")
        driver = webdriver.Chrome()
        driver.get(url)
        time.sleep(2)
        driver.maximize_window()  # 最大化窗口
        driver.find_element_by_class_name("btn_tj").click()
        time.sleep(10)
        #这里是利用模拟浏览器进行自动化点击进行身份验证，全自动只适用于第一级反爬
        #若遇到其他级别的需要人工来验证
        now_url = driver.current_url
        html = reqsuesturl(now_url)
        soup = BeautifulSoup(html, 'lxml')
        driver.quit()
        message1 = (soup.find(class_='c_333 f20 strongbox').string).replace(' ', '_')
    title1 = '大致信息'
    message2_1 = soup.find(class_='f36 strongbox').string
    mark1 = int(message2_1) #这里的标记是指价格，用来分配文件
    message2_2 = str(soup.find(class_='c_ff552e')).split('>')[3][-9:-6]
    message2_3 = soup.find(class_='instructions').string
    message2 = message2_1 + message2_2 + '   ' + message2_3
    # 这里是价格的提取
    message3 = []
    message3_0 = soup.find(class_='f14')
    message3_1 = message3_0.find_all(name='li')
    for i in range(len(message3_1)):
        str1 = str(message3_1[i]).split('</span>')[0][-5:]
        if i == 0:
            str2 = str(message3_1[i]).split('</span>')[-2][6:]
        elif i == 1 or i == 2:
            str2 = str(message3_1[i]).split('>')[4][:-6].replace(' ', '')
        elif i == 3:
            str2 = (message3_1[i].find(class_='c_333 ah').string).replace(' ', '')
            if message3_1[i].find(class_='c_0091d7 ah'):
                str3 = (message3_1[i].find(class_='c_0091d7 ah').string).replace(' ', '')
                str3 = int(str3.replace('\n', ''))
                str2 += '(在租' + str(str3) + '套)'
                mark2 = 1 #这里的标记是指该房是否为小区房
            else:
                mark2 = 0
        elif i == 4:
            list1 = message3_1[i].find_all(class_='c_333 ah')
            str2 = ''
            for i in range(len(list1)):
                x = list1[i].string
                if i is not len(list1) - 1:
                    str2 += x + ' '
                else:
                    str2 += x
        else:
            str2 = (message3_1[i].find(class_='dz').string).replace(' ', '')
            str2 = str2.replace('\n', '')
            str2 = str2[1:]
        strend = str1 + str2
        message3.append(strend)
    # 这是价格下面的信息
    title2 = '房源详情'
    if soup.find(name='div', class_='fang-detail'):
        message4_1 = '卧室设施:'
        message4_2 = '公共设施:'
        x = soup.find_all(class_='house-disposal')
        y = x[0].find_all(name='li')
        for i in range(len(y)):
            if 'no-config' in str(y[i]):
                pass
            else:
                z = str(y[i]).split('</i>')[1][:-5]
                if i is not len(y) - 1:
                    message4_1 += z + ' '
                else:
                    message4_1 += z
        y = x[1].find_all(name='li')
        for i in range(len(y)):
            if 'no-config' in str(y[i]):
                pass
            else:
                z = str(y[i]).split('</i>')[1][:-5]
                if i is not len(y) - 1:
                    message4_2 += z + ' '
                else:
                    message4_2 += z
        message4 = message4_1 + '\n' + message4_2

    else:
        message4 = ''
        x = soup.find(class_='house-disposal')
        y = x.find_all(name='li')
        for i in range(len(y)):
            if 'no-config' in str(y[i]):
                pass
            else:
                z = str(y[i]).split('</i>')[1][:-5]
                if i is not len(y) - 1:
                    message4 += z + ' '
                else:
                    message4 += z
    # 这是房源详情里面的的第一部分
    message5_0 = soup.find(class_='introduce-item')
    message5_1 = message5_0.find_all(class_='a1')
    message5_2 = message5_0.find_all(class_='a2')
    message5 = []
    for i in range(len(message5_1)):
        x = message5_1[i].string + ':'
        if '<em>' in str(message5_2[i]):
            y = ''
            z = re.findall(r'<em>.*?</em>', str(message5_2[i]))
            for j in range(len(z)):
                if j is not len(z) - 1:
                    y += z[j].split('>')[1][:-4] + ' '
                else:
                    y += z[j].split('>')[1][:-4]
        else:
            y = str(message5_2[i])[17:-7].replace('<br/>', '')
        message5.append(x + y)
    # 这是房源详情里面的的第二部分
    message6_1 = soup.find(class_='house-pic-list')
    message6 = []
    list1 = re.findall(r'lazy_src=".*?"', str(message6_1))
    for i in list1:
        x = i.split('"')[1]
        message6.append(x)
    # 这是房源详情里面的的第三部分
    title3 = '小区详情'
    if mark2:
        message7_1 = str(soup.find(class_='c_333 ah rjj')).split('>')[1].replace('\n', '')
        message7_1 = (message7_1.replace(' ', '')[:-14]).replace('\r', '')
        message7_2 = soup.find(class_='addr c_555 f14').string
        message7_3 = soup.find(class_='trend c_0091d7 f14 pr ah')
        message7_3 = '房租走势：' + str(message7_3['href'])
        message7_4 = soup.find(class_='c_333 f24 lh50').string
        message7_4 += '在租房源'
        message7 = [message7_1, message7_2, message7_3, message7_4]
        # 这是小区详情里面的的第一部分
        message8 = soup.find(class_='district-info-list c_333 f14 lh28')
        message8_1 = message8.find_all(class_='c_888 mr_15')
        message8_2 = message8.find_all(name='span', class_='')
        message8 = []
        for i in range(len(message8_2)):
            if i is not len(message8_2) - 1:
                str1 = message8_1[i].string
                str2 = message8_2[i].string
            else:
                str1 = message8_1[i].string
                str3 = re.findall(r'>.*?</a>', str(message8_2[i]))
                str2 = ''
                for i in range(len(str3)):
                    if i is not len(str3) - 1:
                        str2 += str3[i][1:-4] + '/'
                    else:
                        str2 += str3[i][1:-4]
            message8.append(str1 + str2)
    listall = [message1,message2,message3,message4,message5,message6,message7,message8,mark1,mark2]

    return listall
    # 这是小区详情里面的第二部分

def ospath(url,listall):
    if listall[-2] < 1000:
        if os.path.exists('./1000以下'):
            os.chdir(r'./1000以下')
            wordwrite(listall,url)
        else:
            os.mkdir('1000以下')
            os.chdir(r'./1000以下')
            wordwrite(listall,url)
    elif 1000 <= listall[-2] < 2000:
        if os.path.exists('./1000以上2000以下'):
            os.chdir(r'./1000以上2000以下')
            wordwrite(listall,url)
        else:
            os.mkdir('1000以上2000以下')
            os.chdir(r'./1000以上2000以下')
            wordwrite(listall,url)
    else:
        if os.path.exists('./2000以上'):
            os.chdir(r'./2000以上')
            wordwrite(listall,url)
        else:
            os.mkdir('./2000以上')
            os.chdir(r'./2000以上')
            wordwrite(listall,url)

def wordwrite(listall,url):# 开始写入word
    document = docx.Document()
    document.add_heading(listall[0], 0)
    document.add_paragraph(listall[1])
    for i in listall[2]:
        document.add_paragraph(i)
    document.add_paragraph(listall[3])
    for i in listall[4]:
        document.add_paragraph(i)
    for i in range(len(listall[5])):
        r = requests.get(listall[5][i])
        name1 = './' + str(i + 1) + '.png'
        name2 = str(i + 1) + '.png'
        with open(name1, "wb")as f:  # wb是写二进制
            f.write(r.content)
        f = Image.open(name1)
        f.save(name1)
        f.close()
        document.add_picture(name2,width=shared.Inches(5))  # 向文档里添加图片
        os.remove(name2)  # 删除保存在本地的图片
    if listall[-1]:
        for i in listall[6]:
            document.add_paragraph(i)
        for i in listall[7]:
            document.add_paragraph(i)
    driver = webdriver.Chrome()
    driver.get(url)
    driver.maximize_window()
    try:
        driver.find_element_by_link_text("地图详情").click()
        driver.switch_to_window(driver.window_handles[1])
    except:
        driver.find_element_by_class_name('defraud-know').click()
        driver.find_element_by_link_text("地图详情").click()
        driver.switch_to_window(driver.window_handles[1])
    time.sleep(10)
    driver.get_screenshot_as_file("./ditu.png")
    driver.quit()
    document.add_picture('ditu.png',width=shared.Inches(7))
    os.remove('ditu.png')
    name_end = listall[0] + '.docx'
    document.save(name_end)

def main():
    url1 = 'https://fy.58.com/chuzu/pn'
    url2 = '/?key=%E9%98%9C%E9%98%B3%E7%A7%9F%E6%88%BF&final=1&PGTID=0d3090a7-0091-513b-f10d-8f814462c124&ClickID=3'
    number = int(input('请输入想查询的网页数：'))
    urllist = []
    for i in range(1,number+1):
        url = url1 + str(i) + url2
        firsturl(url,urllist)
    for i in urllist:
        print(i)
        listall = secondurl(i)
        ospath(i,listall)
        os.chdir(os.path.pardir)
        print(os.getcwd())
main()